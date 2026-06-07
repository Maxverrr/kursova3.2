from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


DOCX = Path("KVR_Burunduk_Garage_full_report.docx")


MAIN_HEADINGS = {
    "1 Загальний розділ": "ЗАГАЛЬНИЙ РОЗДІЛ",
    "2 Розробка технічного та робочого проєкту": "РОЗРОБКА ТЕХНІЧНОГО ТА РОБОЧОГО ПРОЄКТУ",
    "3 Спеціальний розділ": "СПЕЦІАЛЬНИЙ РОЗДІЛ",
    "4. ЕКОНОМІЧНИЙ РОЗДІЛ": "ЕКОНОМІЧНИЙ РОЗДІЛ",
    "4 Економічний розділ": "ЕКОНОМІЧНИЙ РОЗДІЛ",
    "5 Охорона праці, техніка безпеки та екологічні вимоги": "ОХОРОНА ПРАЦІ, ТЕХНІКА БЕЗПЕКИ ТА ЕКОЛОГІЧНІ ВИМОГИ",
}

UNNUMBERED_H1 = {
    "Зміст": "ЗМІСТ",
    "Вступ": "ВСТУП",
    "Висновки": "ВИСНОВКИ",
    "Перелік посилань": "ПЕРЕЛІК ПОСИЛАНЬ",
    "ДОДАТКИ": "ДОДАТКИ",
    "АНОТАЦІЯ": "АНОТАЦІЯ",
}

TOC_ENTRIES = [
    ("ВСТУП", "toc_intro", 1),
    ("1\tЗАГАЛЬНИЙ РОЗДІЛ", "toc_sec1", 1),
    ("1.1\tАналітичний огляд існуючих рішень", "toc_sec1_1", 2),
    ("1.2\tТехнічне завдання", "toc_sec1_2", 2),
    ("1.2.1\tНайменування та область застосування", "toc_sec1_2_1", 3),
    ("1.2.2\tПризначення розробки", "toc_sec1_2_2", 3),
    ("1.2.3\tВимоги до функціоналу web-сайту", "toc_sec1_2_3", 3),
    ("1.2.4\tВимоги до програмної документації", "toc_sec1_2_4", 3),
    ("1.2.5\tТехніко-економічні показники", "toc_sec1_2_5", 3),
    ("1.2.6\tСтадії та етапи розробки", "toc_sec1_2_6", 3),
    ("1.2.7\tПорядок тестування та прийому", "toc_sec1_2_7", 3),
    ("2\tРОЗРОБКА ТЕХНІЧНОГО ТА РОБОЧОГО ПРОЄКТУ", "toc_sec2", 1),
    ("2.1\tРозробка структури сайту і web-сторінок", "toc_sec2_1", 2),
    ("2.2\tСтворення та верстка сторінок сайту", "toc_sec2_2", 2),
    ("2.3\tРозробка структури бази даних сайту", "toc_sec2_3", 2),
    ("2.4\tПрограмування сайту", "toc_sec2_4", 2),
    ("2.4.1\tНаписання клієнтської частини", "toc_sec2_4_1", 3),
    ("2.4.2\tНаписання admin- частини", "toc_sec2_4_2", 3),
    ("2.5\tТестування web- сайту", "toc_sec2_5", 2),
    ("3\tСПЕЦІАЛЬНИЙ РОЗДІЛ", "toc_sec3", 1),
    ("3.1\tІнструкція з розміщення сайту в Інтернеті", "toc_sec3_1", 2),
    ("3.2\tІнструкція з обслуговування та наповнення сайту", "toc_sec3_2", 2),
    ("3.3\tІнструкція з популяризації та підтримки сайту", "toc_sec3_3", 2),
    ("4\tЕКОНОМІЧНИЙ РОЗДІЛ", "toc_sec4", 1),
    ("5\tОХОРОНА ПРАЦІ, ТЕХНІКА БЕЗПЕКИ ТА ЕКОЛОГІЧНІ ВИМОГИ", "toc_sec5", 1),
    ("ВИСНОВКИ", "toc_conclusions", 1),
    ("ПЕРЕЛІК ПОСИЛАНЬ", "toc_refs", 1),
]

BOOKMARK_BY_HEADING = {
    "ВСТУП": "toc_intro",
    "ЗАГАЛЬНИЙ РОЗДІЛ": "toc_sec1",
    "Аналітичний огляд існуючих рішень": "toc_sec1_1",
    "Технічне завдання": "toc_sec1_2",
    "Найменування та область застосування": "toc_sec1_2_1",
    "Призначення розробки": "toc_sec1_2_2",
    "Вимоги до функціоналу web-сайту": "toc_sec1_2_3",
    "Вимоги до програмної документації": "toc_sec1_2_4",
    "Техніко-економічні показники": "toc_sec1_2_5",
    "Стадії та етапи розробки": "toc_sec1_2_6",
    "Порядок тестування та прийому": "toc_sec1_2_7",
    "РОЗРОБКА ТЕХНІЧНОГО ТА РОБОЧОГО ПРОЄКТУ": "toc_sec2",
    "Розробка структури сайту і web-сторінок": "toc_sec2_1",
    "Створення та верстка сторінок сайту": "toc_sec2_2",
    "Розробка структури бази даних сайту": "toc_sec2_3",
    "Програмування сайту": "toc_sec2_4",
    "Написання клієнтської частини": "toc_sec2_4_1",
    "Написання admin- частини": "toc_sec2_4_2",
    "Тестування web- сайту": "toc_sec2_5",
    "СПЕЦІАЛЬНИЙ РОЗДІЛ": "toc_sec3",
    "Інструкція з розміщення сайту в Інтернеті": "toc_sec3_1",
    "Інструкція з обслуговування та наповнення сайту": "toc_sec3_2",
    "Інструкція з популяризації та підтримки сайту": "toc_sec3_3",
    "ЕКОНОМІЧНИЙ РОЗДІЛ": "toc_sec4",
    "ОХОРОНА ПРАЦІ, ТЕХНІКА БЕЗПЕКИ ТА ЕКОЛОГІЧНІ ВИМОГИ": "toc_sec5",
    "ВИСНОВКИ": "toc_conclusions",
    "ПЕРЕЛІК ПОСИЛАНЬ": "toc_refs",
}


def run_font(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def clear_runs(paragraph, text):
    paragraph.text = text
    for run in paragraph.runs:
        run_font(run, bold=paragraph.style.name.startswith("Heading"))


def set_num_pr(container, num_id, ilvl=None):
    ppr = container._element.get_or_add_pPr() if hasattr(container, "_element") else container._p.get_or_add_pPr()
    if ppr.numPr is not None:
        ppr.remove(ppr.numPr)
    num_pr = OxmlElement("w:numPr")
    if ilvl is not None:
        ilvl_el = OxmlElement("w:ilvl")
        ilvl_el.set(qn("w:val"), str(ilvl))
        num_pr.append(ilvl_el)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_el)
    ppr.append(num_pr)


def remove_num_pr(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    if ppr.numPr is not None:
        ppr.remove(ppr.numPr)
    num_pr = OxmlElement("w:numPr")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), "0")
    num_pr.append(num_id_el)
    ppr.append(num_pr)


def add_bookmark(paragraph, name, bid):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_hyperlink(paragraph, text, anchor):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rpr.append(color)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_complex_field(paragraph, instruction, placeholder=""):
    r = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r._r.append(begin)
    r = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " " + instruction + " "
    r._r.append(instr)
    r = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r._r.append(sep)
    if placeholder:
        rr = paragraph.add_run(placeholder)
        run_font(rr)
    r = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(end)


def add_field_begin_and_separator(paragraph, instruction):
    r = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r._r.append(begin)
    r = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " " + instruction + " "
    r._r.append(instr)
    r = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r._r.append(sep)


def insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def configure_toc_style(doc, name, left_cm):
    style = doc.styles[name]
    pf = style.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(left_cm)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(14)


def configure_styles(doc, heading_num_id):
    for style_name, ilvl in [("Heading 1", 0), ("Heading 2", 1), ("Heading 3", 2)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.font.bold = True
        set_num_pr(style, heading_num_id, None if ilvl == 0 else ilvl)
        pf = style.paragraph_format
        pf.first_line_indent = Cm(1.25 if ilvl else 0)
        pf.left_indent = Cm(0)
        pf.space_before = Pt(12 if ilvl else 12)
        pf.space_after = Pt(6 if ilvl else 6)
        pf.line_spacing = 1.5
    configure_toc_style(doc, "toc 1", 0.5)
    configure_toc_style(doc, "toc 2", 0.5)
    configure_toc_style(doc, "toc 3", 1.0)


def normalize_heading_text(doc):
    for p in doc.paragraphs:
        if not p.style.name.startswith("Heading"):
            continue
        text = p.text.strip()
        if not text:
            continue
        if text in MAIN_HEADINGS:
            clear_runs(p, MAIN_HEADINGS[text])
            continue
        if text in UNNUMBERED_H1:
            clear_runs(p, UNNUMBERED_H1[text])
            remove_num_pr(p)
            continue
        if text.startswith("Додаток"):
            remove_num_pr(p)
            continue
        m = re.match(r"^\d+(?:\.\d+)*\.?\s+(.+)$", text)
        if m and p.style.name in ("Heading 2", "Heading 3"):
            clear_runs(p, m.group(1).strip())
        elif m and p.style.name == "Heading 1":
            clear_runs(p, m.group(1).strip().upper())
        if p.text.strip() in UNNUMBERED_H1:
            remove_num_pr(p)


def add_heading_bookmarks(doc):
    bid = 900
    seen = set()
    for p in doc.paragraphs:
        text = p.text.strip()
        if p.style.name.startswith("Heading") and text in BOOKMARK_BY_HEADING:
            name = BOOKMARK_BY_HEADING[text]
            if name not in seen:
                add_bookmark(p, name, bid)
                bid += 1
                seen.add(name)


def rebuild_toc(doc):
    toc_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().upper() == "ЗМІСТ")
    next_idx = next(
        i for i in range(toc_idx + 1, len(doc.paragraphs))
        if doc.paragraphs[i].style.name.startswith("Heading") and doc.paragraphs[i].text.strip().upper() == "ВСТУП"
    )
    for p in list(doc.paragraphs[toc_idx + 1:next_idx]):
        delete_paragraph(p)
    toc_heading = doc.paragraphs[toc_idx]
    toc_heading.style = doc.styles["TOC Heading"] if "TOC Heading" in [s.name for s in doc.styles] else doc.styles["Heading 1"]
    clear_runs(toc_heading, "ЗМІСТ")
    remove_num_pr(toc_heading)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Word TOC field like in the sample; entries below are the current field result.
    anchor = toc_heading
    start = insert_paragraph_after(anchor)
    start.paragraph_format.space_after = Pt(0)
    add_field_begin_and_separator(start, 'TOC \\o "1-3" \\h \\z \\u')
    anchor = start
    for label, bookmark, level in TOC_ENTRIES:
        p = insert_paragraph_after(anchor)
        p.style = doc.styles[f"toc {level}"]
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16.0), alignment=2, leader=1)
        parts = label.split("\t", 1)
        if len(parts) == 2:
            add_hyperlink(p, parts[0], bookmark)
            p.add_run("\t")
            add_hyperlink(p, parts[1], bookmark)
        else:
            add_hyperlink(p, label, bookmark)
        p.add_run("\t")
        add_complex_field(p, f"PAGEREF {bookmark} \\h", "1")
        anchor = p
    end = insert_paragraph_after(anchor)
    r = end.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r._r.append(fld_end)


def patch_numbering_and_settings(docx_path, heading_num_id=99, abstract_id=99):
    with zipfile.ZipFile(docx_path, "a") as z:
        nums = z.read("word/numbering.xml").decode("utf-8")
        if f'w:numId="{heading_num_id}"' not in nums:
            abstract = f'''
<w:abstractNum w:abstractNumId="{abstract_id}">
  <w:multiLevelType w:val="multilevel"/>
  <w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:pStyle w:val="11"/><w:lvlText w:val="%1"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="432" w:hanging="432"/></w:pPr></w:lvl>
  <w:lvl w:ilvl="1"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:pStyle w:val="20"/><w:lvlText w:val="%1.%2"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="576" w:hanging="576"/></w:pPr></w:lvl>
  <w:lvl w:ilvl="2"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:pStyle w:val="3"/><w:lvlText w:val="%1.%2.%3"/><w:lvlJc w:val="left"/><w:pPr><w:ind w:left="1571" w:hanging="720"/></w:pPr></w:lvl>
</w:abstractNum>
<w:num w:numId="{heading_num_id}"><w:abstractNumId w:val="{abstract_id}"/></w:num>
'''
            nums = nums.replace("</w:numbering>", abstract + "</w:numbering>")
            z.writestr("word/numbering.xml", nums)
        settings = z.read("word/settings.xml").decode("utf-8")
        if "w:updateFields" not in settings:
            settings = settings.replace("</w:settings>", '<w:updateFields w:val="true"/></w:settings>')
            z.writestr("word/settings.xml", settings)


def normalize_zip(docx_path):
    import os
    import shutil
    import tempfile

    fd, tmp = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    tmp = Path(tmp)
    with zipfile.ZipFile(docx_path, "r") as zin:
        names = zin.namelist()
        keep = {name: i for i, name in enumerate(names)}
        with zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for i, name in enumerate(names):
                if keep[name] == i:
                    zout.writestr(name, zin.read(name))
    shutil.move(str(tmp), str(docx_path))


def main():
    heading_num_id = 99
    doc = Document(DOCX)
    # Keep TOC away from the bottom stamp/footer area.
    if doc.sections:
        doc.sections[0].bottom_margin = Cm(3.0)
        doc.sections[0].footer_distance = Cm(0.7)
    configure_styles(doc, heading_num_id)
    normalize_heading_text(doc)
    add_heading_bookmarks(doc)
    rebuild_toc(doc)
    doc.save(DOCX)
    patch_numbering_and_settings(DOCX, heading_num_id=heading_num_id, abstract_id=99)
    normalize_zip(DOCX)


if __name__ == "__main__":
    main()
