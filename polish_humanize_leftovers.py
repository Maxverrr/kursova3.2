from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX = Path("KVR_Burunduk_Garage_full_report.docx")


def run_font(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def format_body(paragraph):
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        run_font(run)


def first_heading_index(doc, prefix):
    return next(
        i for i, p in enumerate(doc.paragraphs)
        if p.style.name.startswith("Heading") and p.text.strip().startswith(prefix)
    )


def main():
    doc = Document(DOCX)
    idx4 = first_heading_index(doc, "4")
    after4_snapshot = [(p.style.name, p.text) for p in doc.paragraphs[idx4:]]

    replacements = [
        (
            "Для «Burunduk Garage» доцільним є саме індивідуальний вебзастосунок, оскільки він поєднує каталог автомобілів, особистий кабінет користувача, адміністративний облік бронювань, статистику та роботу із зображеннями. Такий підхід забезпечує контроль над функціональністю і дозволяє реалізувати деталі, які важливі саме для конкретного автопарку.",
            "Для «Burunduk Garage» доцільним є саме індивідуальний вебзастосунок, оскільки він поєднує каталог автомобілів, особистий кабінет користувача, адміністративний облік бронювань, статистику та роботу із зображеннями. У готових сервісах подібні дрібні налаштування часто довелося б обходити сторонніми рішеннями, а тут їх можна закласти одразу в логіку проєкту.",
        ),
        (
            "Для сервісів оренди автомобілів важливо не лише показати каталог машин, а й забезпечити зв'язок між наявністю автомобіля, датами оренди, клієнтськими даними та діями адміністратора. Багато невеликих автопарків починають роботу з таблиць або месенджерів, однак такий підхід швидко стає незручним, коли збільшується кількість клієнтів і змінюються умови бронювань.",
            "Сервісу оренди автомобілів недостатньо просто показати каталог машин: потрібно пов'язати наявність автомобіля, дати оренди, клієнтські дані та дії працівника. Багато невеликих автопарків починають роботу з таблиць або месенджерів, але зі збільшенням кількості клієнтів такий облік швидко стає незручним.",
        ),
        (
            "Важливо, щоб сторінка автопарку відкривалася швидко і плавно прокручувалася.",
            "Сторінка автопарку має відкриватися швидко і плавно прокручуватися.",
        ),
        (
            "Для сервісу оренди важливо, щоб сторінка автопарку відкривалася швидко і плавно прокручувалася.",
            "Для сервісу оренди сторінка автопарку має відкриватися швидко і плавно прокручуватися.",
        ),
        (
            "Для адміністратора це дає загальне уявлення про стан автопарку без ручних підрахунків.",
            "У результаті стан автопарку видно без ручних підрахунків у таблицях.",
        ),
        (
            "У цьому проєкті потрібно було врахувати власну логіку: редагування та скасування для користувача тільки за два дні до початку оренди, окремий пошук у бронюваннях і підтримку електромобілів без зміни серверної моделі.",
            "У цьому проєкті потрібно було врахувати власну логіку: редагування та скасування для користувача тільки за два дні до початку оренди, окремий пошук у бронюваннях і гнучкі підписи характеристик без зміни серверної моделі.",
        ),
    ]

    for paragraph in doc.paragraphs[:idx4]:
        text = paragraph.text
        changed = False
        for old, new in replacements:
            if old in text:
                text = text.replace(old, new)
                changed = True
        if changed:
            paragraph.text = text
            format_body(paragraph)

    new_idx4 = first_heading_index(doc, "4")
    if [(p.style.name, p.text) for p in doc.paragraphs[new_idx4:]] != after4_snapshot:
        raise RuntimeError("Unexpected change after section 4")

    doc.save(DOCX)


if __name__ == "__main__":
    main()
