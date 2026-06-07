from docx import Document


PHRASES = [
    "Для адміністратора",
    "Для користувача",
    "Таким чином",
    "Окрема увага",
    "Важливо, щоб",
    "Такий підхід",
    "електромоб",
    "режим гостя",
    "гість",
]

doc = Document("KVR_Burunduk_Garage_full_report.docx")
idx4 = next(
    (i for i, p in enumerate(doc.paragraphs) if p.style.name.startswith("Heading") and p.text.strip().startswith("4")),
    len(doc.paragraphs),
)
print("idx4", idx4)
for phrase in PHRASES:
    hits = []
    for i, p in enumerate(doc.paragraphs[:idx4]):
        if phrase.lower() in p.text.lower():
            hits.append((i, p.text[:220]))
    print("\n" + phrase, len(hits))
    for i, text in hits[:20]:
        print(i, text)
