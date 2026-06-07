from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


DOCX = Path("KVR_Burunduk_Garage_full_report.docx")


def run_font(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def format_body(paragraph):
    paragraph.style = paragraph.part.document.styles["Normal"]
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        run_font(run)


def insert_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_paragraph.add_run(text)
    format_body(new_paragraph)
    return new_paragraph


def headings(doc):
    return [
        (i, p.text.strip(), p)
        for i, p in enumerate(doc.paragraphs)
        if p.style.name.startswith("Heading") and p.text.strip()
    ]


def anchor_for(doc, heading_text):
    hs = headings(doc)
    start = next(i for i, text, _ in hs if text == heading_text)
    next_heading = len(doc.paragraphs)
    for i, p in enumerate(doc.paragraphs[start + 1 :], start + 1):
        if p.style.name.startswith("Heading") and p.text.strip():
            next_heading = i
            break
    idx = next_heading - 1
    while idx > start and not doc.paragraphs[idx].text.strip():
        idx -= 1
    return doc.paragraphs[idx]


extra = {
    "3.1 Інструкція з розміщення сайту в Інтернеті": [
        "Перед розгортанням frontend потрібно переконатися, що локальна збірка виконується без помилок. Для цього у папці клієнтської частини запускається встановлення залежностей і команда збірки. Якщо Vite успішно формує папку dist, проєкт готовий до публікації на Vercel.",
        "У Vercel доцільно підключати саме той репозиторій, у якому зберігається актуальна версія frontend. Після кожного push у вибрану гілку Vercel може автоматично запускати новий deployment. Це зручно, оскільки зміни в інтерфейсі автопарку або сторінках бронювань одразу потрапляють у тестове чи робоче середовище.",
        "Для backend на Render важливо перевірити, що сервер слухає порт із середовища, а не тільки фіксоване локальне значення. На Render порт передається платформою, тому сервер Express повинен використовувати process.env.PORT з резервним значенням для локального запуску.",
        "MongoDB краще використовувати як окремий хмарний сервіс, наприклад MongoDB Atlas. У Render зберігається тільки рядок підключення, а сама база даних залишається незалежною від сервера. Це спрощує перенесення backend і зменшує ризик втрати даних під час оновлення сервісу.",
        "Після першого розгортання потрібно перевірити CORS у реальних умовах. Якщо frontend відкривається, але не завантажує автомобілі, причина часто полягає у неправильній адресі API або забороні запитів із домену Vercel. У такому випадку необхідно оновити список дозволених origin на backend.",
        "Також потрібно перевірити роботу ImageKit після розгортання. Локально URL може формуватися правильно, але на сервері мають бути задані всі ключі ImageKit. Якщо хоча б один ключ відсутній, додавання або оновлення фото автомобіля може завершуватися помилкою.",
    ],
    "3.2 Інструкція з обслуговування та наповнення сайту": [
        "Після розміщення сайту адміністратор повинен вести автопарк як актуальний електронний каталог. Якщо автомобіль змінює ціну, статус або фото, ці дані потрібно оновити в admin-панелі до того, як клієнт оформить нове бронювання.",
        "Під час наповнення сайту варто дотримуватися єдиного стилю описів. Назви автомобілів, характеристики, витрата палива, дані про електромобілі та ціна повинні подаватися однаково для всіх карток. Це робить автопарк зручним для порівняння.",
        "Обслуговування бронювань передбачає регулярну перевірку майбутніх оренд. Адміністратор може використовувати фільтр за датами, щоб бачити найближчі видачі автомобілів, або фільтр за завершеними орендами, щоб аналізувати історію роботи сервісу.",
        "Якщо користувач звертається з проханням змінити бронювання, адміністратор може знайти його за іменем клієнта або автомобілем. Після цього перевіряється новий період оренди і, якщо він не конфліктує з іншими бронюваннями, дані оновлюються в системі.",
        "Для стабільної роботи сайту потрібно періодично перевіряти логи Render. У логах можуть бути помилки підключення до бази даних, неправильні токени, некоректні запити або проблеми з ImageKit. Регулярний перегляд логів допомагає знаходити проблеми до того, як їх помітять клієнти.",
    ],
    "3.3 Інструкція з популяризації та підтримки сайту": [
        "Під час популяризації Burunduk Garage доцільно використовувати прямі посилання на сторінку автопарку. Користувач повинен одразу бачити доступні автомобілі, а не шукати каталог через кілька проміжних сторінок. Це скорочує шлях від реклами до бронювання.",
        "Для соціальних мереж можна готувати публікації з окремими автомобілями, вказуючи основні характеристики та посилання на детальну сторінку. Якщо фото в картках мають однаковий стиль завдяки ImageKit, візуальна подача бренду виглядає цілісніше.",
        "Підтримка сайту включає не тільки технічні оновлення, а й аналіз поведінки клієнтів. Якщо певні автомобілі часто переглядають, але рідко бронюють, варто перевірити ціну, фото, опис або доступність дат. Такі спостереження можуть впливати на розвиток автопарку.",
        "У майбутньому сайт можна доповнити сторінками з умовами оренди, відповідями на часті запитання, блоком акцій та окремими посадковими сторінками для популярних класів автомобілів. Це підвищить інформативність сайту і може покращити його позиції у пошуку.",
        "Важливо, щоб підтримка не порушувала стабільність уже реалізованих функцій. Перед додаванням нових можливостей потрібно перевіряти авторизацію, бронювання, фільтри, роботу admin-панелі та адаптивність. Для сервісу оренди навіть невелика помилка в датах може мати практичні наслідки.",
    ],
}


def main():
    doc = Document(DOCX)
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading") and not paragraph.text.strip():
            paragraph.style = doc.styles["Normal"]
            paragraph.paragraph_format.first_line_indent = Cm(0)

    for heading_text, paragraphs in extra.items():
        anchor = anchor_for(doc, heading_text)
        for text in paragraphs:
            anchor = insert_after(anchor, text)

    doc.save(DOCX)


if __name__ == "__main__":
    main()
