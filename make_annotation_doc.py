from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


sample = Path("input_annotation_sample.docx")
out = Path("Anotatsiya_Burunduk_Garage.docx")

doc = Document(str(sample))

texts = [
    "АНОТАЦІЯ",
    "",
    "",
    "Тема кваліфікаційної роботи: Розробка вебзастосунку для автоматизації управління автопарком «Burunduk Garage».",
    "Метою кваліфікаційної роботи є розробка web-застосунку для автоматизації управління автопарком «Burunduk Garage», що забезпечує ведення каталогу автомобілів, оформлення та редагування бронювань, облік користувачів, роботу адміністративної панелі, перегляд статистики та зручну взаємодію клієнта із сервісом оренди автомобілів.",
    "Пояснювальна записка складається з п’яти розділів.",
    "У загальній частині описуються аналітичний огляд існуючих рішень у сфері автоматизації автопрокату та аналіз технічного завдання на розробку web-сайту «Burunduk Garage».",
    "У другому розділі представлено процес створення програмного продукту, опис та обґрунтування вибору структури сайту і web-сторінок, створення та верстку сторінок, розробку структури бази даних сайту, програмування клієнтської та admin-частини, а також тестування web-сайту.",
    "В спеціальній частині описані інструкція з розміщення сайту в Інтернеті, інструкція з обслуговування та наповнення сайту, а також інструкція з популяризації та підтримки web-сайту після впровадження.",
    "Розрахунок вартості розробки та економічної ефективності приведено в економічній частині.",
    "Основні питання охорони праці, техніки безпеки та екологічні вимоги розглянуто в п’ятому розділі.",
    "Обсяг пояснювальної записки _____ сторінок.",
    "До складу кваліфікаційної роботи входить графічна частина, яка складається зі структурної схеми web-сайту, схеми бази даних, блок-схеми алгоритму бронювання автомобіля, схеми взаємодії клієнтської частини, серверної частини, MongoDB та ImageKit, техніко-економічних показників і лістингів основних програмних модулів, що виконані на окремих аркушах формату А1.",
    "",
]

while len(doc.paragraphs) < len(texts):
    doc.add_paragraph()

for par, text in zip(doc.paragraphs, texts):
    par.text = text
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.space_after = Pt(0)
    par.paragraph_format.first_line_indent = Cm(1.25)
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text == "АНОТАЦІЯ":
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.first_line_indent = Cm(0)
    if not text:
        par.paragraph_format.first_line_indent = Cm(0)
    for run in par.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)
        run.bold = text == "АНОТАЦІЯ"

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.0)

doc.save(out)
print(out.resolve())
